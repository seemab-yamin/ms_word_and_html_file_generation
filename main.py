from docx import Document
from os import path as Path

if __name__ == '__main__':
    #input file name
    inputFilePath = 'LoremIpsum.txt'

    
    #create new document
    document = Document()

    #read input .txt file
    with open(inputFilePath, 'r') as file:
        paragraphs = file.readlines()
    
    paragraphs = [ para.strip() for para in paragraphs if para!='\n' ]
    formattedParagraphs = []

    htmlText = ""
    for para in paragraphs:
        paraObj = document.add_paragraph()
        htmlText += "<p>"
        for word in para.split(' '):
            endsWith = ' '
            if '.' in word:
                word = word.replace('.', '')
                endsWith = '. '
            elif ',' in word:
                word = word.replace(',', '')
                endsWith = ', '

            if len(word)>1:
                halfIndex = (len(word) // 2) + 1
                firstHalf = word[:halfIndex]
                secondHalf = word[halfIndex : ]
                
                boldRun = paraObj.add_run(firstHalf)
                boldRun.bold = True
                normalRun = paraObj.add_run(secondHalf + endsWith)

                htmlText += "<strong>" + firstHalf + "</strong>"
                htmlText += secondHalf + endsWith
            else:
                singleLetterRun = paraObj.add_run(word + endsWith)
                singleLetterRun.bold = True
                htmlText += word + endsWith
        htmlText += "</p>"

    #save as new document
    fileName = Path.basename(inputFilePath).split('.')[0]
    outputFileName = fileName + " (output generated)"

    with open(outputFileName + '.html', 'w') as file:
        file.write(htmlText)

    document.save(Path.join(Path.curdir, outputFileName + '.docx'))